from itertools import product
from docx import Document
from docx.shared import Pt, Inches

# Define the character sets
start = ["ن"]
mid = ["ب", "ج", "س", "ص", "ط", "ع", "ف", "ق", "ک", "ل", "م", "ن", "ہ", "ی"]
end = ["ا", "ب", "ج", "د", "ر", "س", "ص", "ط", "ع", "ف", "ق", "ک", "ل", "م", "ن", "و", "ہ", "ی", "ے"]

# Create a new Word document
doc = Document()
section = doc.sections[0]
section.page_height = Inches(11.69)  # A4 height
section.page_width = Inches(8.27)    # A4 width
section.top_margin = Inches(0.5)     # Narrow margins (0.5 inch)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

# Set the document style
style = doc.styles['Normal']
style.font.name = 'Jameel Noori Nastaleeq'
style.font.size = Pt(40)  # Adjust font size if needed
style.paragraph_format.line_spacing = Pt(90)  # Adjust line spacing for better fit

# Create a table with 4 columns
table = doc.add_table(rows=0, cols=4)

# Function to add a word to the table
def add_word_to_table(table, word):
    cells = table.add_row().cells
    for i in range(4):
        if i == 0:
            run = cells[i].paragraphs[0].add_run(word)
            run.font.size = Pt(40)
        else:
            cells[i].text = ''

# Generate all possible 5-character words
words = []
for s in start:
    for mid_combo in product(mid, repeat=3):
        for e in end:
            word = s + "".join(mid_combo) + e
            words.append(word)

# Add words to the table in rows of 4
for i in range(0, len(words), 4):
    row_words = words[i:i + 4]
    row = table.add_row().cells
    for j, word in enumerate(row_words):
        run = row[j].paragraphs[0].add_run(word + ' ' * 5)  # Adding spaces for separation
        run.font.size = Pt(40)

# Save the document
doc.save('urdu_words.docx')
